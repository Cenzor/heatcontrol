from django.core.management.base import BaseCommand
from docx import Document
from django.conf import settings
from django.template import engines
from copy import deepcopy
from django.utils import timezone
from heatcontrol.api.models import MeteringPointData, User, CompanyObject


def hide_column(table, column_index):
    table.columns[column_index].width = 0
    for cell in table.columns[column_index].cells:
        cell.text = ""


class Command(BaseCommand):
    help = 'create doc'


    def handle(self, *args, **options):
        DATA_TABLE_INDEX = 2
        data = [
            {
                "report_type": "CURRENT",
#                 "date": timezone.now(),
                "date_from": timezone.now(),
                "date_to": timezone.now(),
                "company_object": CompanyObject.objects.last(),
                "user": User.objects.last(),
                "rows": list(MeteringPointData.objects.all().order_by("-id")[:3]),
            },
            {
                "report_type": "HOURLY",
                "date": timezone.now(),
#                 "date_from": timezone.now(),
#                 "date_to": timezone.now(),
                "company_object": CompanyObject.objects.last(),
                "user": User.objects.last(),
                "rows": list(MeteringPointData.objects.all().order_by("-id")[:2]),
            },
            {
                "report_type": "DAILY",
                "date": timezone.now(),
                "date_from": timezone.now(),
                "date_to": timezone.now(),
                "company_object": CompanyObject.objects.last(),
                "user": User.objects.last(),
                "rows": list(MeteringPointData.objects.all().order_by("-id")[:1]),
            },
        ]
        headers = "t1,t2,dt,V1,M1,V2,M2,P1,P2,Qо,BНP,BOC,НС"
        header_values = headers.split(",")
        values = []
        for header_value in header_values:
            values.append(header_value.lower())
        header_values = values
        django_engine = engines['django']
        document = Document("%s/templates/report.docx" % settings.BASE_DIR)
        #запомним, сколько изначально таблиц в документе
        tables_amount = len(document.tables)
        for i in range(1, len(data)):
            #добавляем страницы
            document.add_page_break()
            for table_index in range(0, tables_amount):
                template = document.tables[table_index]
                tbl = template._tbl
                # Here we do the copy of the table
                new_tbl = deepcopy(tbl)    
                paragraph = document.add_paragraph()
                # After that, we add the previously copied table
                paragraph._p.addnext(new_tbl)
        tables = document.tables
        table_index = 0
        for context in data:
            for i in range(0, tables_amount):
                table = tables[table_index + i]
                columns = len(table.columns)
                if i == DATA_TABLE_INDEX:
                    #это таблица с данными. там необходимо сделать столько строк, сколько строк с данными и в цикле считать туда все данные
                    #уберем колонки, отсутствующие в данных
                    for column in range(1, columns):
                        value = table.cell(0, column).text
                        value = value.split(",")[0].split("\n")[0].strip().lower()
                        print("header = %s" % value)
                        if not value in header_values:
                            print("hidding %s" % value)
                            hide_column(table, column)
                    if len(context["rows"]) > 1:
                        template_row = table.rows[1]
                        for row_index in range(1, len(context["rows"])):
                            row = table.add_row()
                            tr = template_row._tr
                            content = deepcopy(tr)
                            row._tr = content
                            
#                             for col_index in range(0, len(template_row.cells)):
#                                 template_cell = template_row.cells[col_index]
#                                 template_paragraph = template_cell.paragraphs[0]
#                                 cell = row.cells[col_index]
#                                 cell.text = template_paragraph.text
#                                 paragraph = cell.paragraphs[0]
#                                 paragraph.style = template_paragraph.style
                    for row_index in range(0, len(context["rows"])):
                        context_row = context["rows"][row_index]
                        context["row"] = context_row
                        for column in range(0, columns):
                            cell = table.cell(1 + row_index, column)
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                template = django_engine.from_string("{% load call_method %}" + paragraph.text)
                                rendered_text = template.render(context)
                                if not paragraph.text == rendered_text:
                                    paragraph.text = rendered_text
                #заполняем остальные поля
                rows = len(table.rows)
                for row in range(0, rows):
                    for column in range(0, columns):
                        cell = table.cell(row, column)
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            template = django_engine.from_string("{% load call_method %}" + paragraph.text)
                            rendered_text = template.render(context)
                            if not paragraph.text == rendered_text:
                                paragraph.text = rendered_text
            table_index += tables_amount
        

        

        document.save("report.docx")