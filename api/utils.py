from django.conf import settings
from django.template import loader
from django.template.base import Template
from django.core.mail.message import EmailMessage
import traceback
from django.core import mail
from heatcontrol.api.models import DataEditHistory, CompanyObject, MeteringPoint,\
    MeteringPointData, Department, DepartmentType
from django.utils import dateparse, timezone
import json
from docx import Document
from django.template import engines
from copy import deepcopy
from wkhtmltopdf.views import PDFTemplateResponse


def send_email(to, subject_template, message_template, from_address = None, use_html = False, **kwargs):
    if settings.TESTING:
#         print("skipping email sending in test mode")
        return
    """
    Send email message to address given, using kwargs as a context variable
    """

#     kwargs.update({
#         "site_domain": settings.SITE_DOMAIN
#     })
    if not "subject" in kwargs:
        t = loader.get_template(subject_template)
        subject = t.render(kwargs)
    else:
        t =  Template(subject)
        subject = t.render(kwargs)
    
    # Email subject *must not* contain newlines
    subject = ''.join(subject.splitlines())

    t = loader.get_template(message_template)
    message = t.render(kwargs)

    if not from_address:
        from_address = settings.DEFAULT_FROM_EMAIL

    
    print("from_address = %s" % from_address)

    headers = dict()
    reply_to = from_address
    if "reply_to" in kwargs:
        reply_to = kwargs.get("reply_to")
        headers = {'Reply-To': reply_to}
    
    cc = None
    if "cc" in kwargs:
        cc = kwargs.get("cc")

    if not type(to) is list:
        to = [to]

    print("to = %s" % to)

    try:
        msg = None
        if cc:
            msg = EmailMessage(subject, message, from_address, to, headers=headers, cc=[cc])
        else:
            msg = EmailMessage(subject, message, from_address, to, headers=headers)
        if use_html:
            msg.content_subtype = "html"
        msg.send()    
#         mail.send_mail(subject, message, from_address, [to], headers=headers)
    except Exception as e:
        traceback.print_exc()
#         logger.warning('Email util: failed to deliver message', exc_info=sys.exc_info())
        #notify admin that there was an exception
        try:
            subject = "ERROR while sending: %s" % subject
            message = "The following message was not delivered to %s (exception while sending email). Original message text: %s" % (to, message)
            mail.send_mail(subject, message, from_address, [settings.ADMINS[0][1]])
        except Exception as e:
            traceback.print_exc()


def track_object_updated(action, user, obj, obj_name, new_data = None):
    meta = obj._meta
    model_name = meta.object_name
    if action == "Edit":
        history = dict()
        for field in meta.fields:
            field_name = field.attname
            if field_name in new_data:
                new_value = str(new_data[field_name])
                old_value = str(getattr(obj, field_name))
                if not new_value == old_value:
                    history[field_name] = {
                        "old": old_value,
                        "new": new_value,
                    }
        if history:
            data_edit_history = DataEditHistory(
                action = action,
                user = user,
                model_name = model_name,
                object_id = obj.id,
                object_name = obj_name,
                data = history,
            )
            data_edit_history.save()
    else:
        data_edit_history = DataEditHistory(
            action = action,
            user = user,
            model_name = model_name,
            object_id = obj.id,
            object_name = obj_name,
        )
        data_edit_history.save()


def generate_report(request, for_excel=False, report_user=None, report_task=None):
    report_type = None
    date_from = None
    date_to = None
    object_ids = None
    company_objects = None
    user = None
    if request:
        report_type = request.GET.get("report_type", "DAILY")
        objects = request.GET.get("objects", "")
        date_from = request.GET.get("date_from", None)
        date_to = request.GET.get("date_to", None)
        if date_from:
            date_from = dateparse.parse_datetime(date_from)
        if date_to:
            date_to = dateparse.parse_datetime(date_to)
        object_ids = None
        if objects:
            object_ids = json.loads("[%s]" % objects)
    #         print("object_ids = " , object_ids)
        user = request.user
        permissions = request.user.get_permissions()
        company_objects = CompanyObject.objects.filter(deleted = False)
        if not permissions.status_admin_system:
            company_objects = company_objects.filter(company = user.employer)
        if object_ids:
            company_objects = company_objects.filter(id__in = object_ids)
    #         print("company_objects = ", company_objects)
        object_ids = company_objects.values_list("id", flat = True)
    else:
        report_type = report_task.report_type
        date_from = report_task.date_from
        date_to = report_task.date_to
        object_ids = []
        for object in report_task.company_objects.all():
            object_ids.append(object.id)
        company_objects = report_task.company_objects.all()

    if for_excel:
        result = {'company_objects': []}
        for company_ob in company_objects:
            rec = {'company_object': company_ob,
                   'metering_points': []}
            metering_points = MeteringPoint.objects.filter(company_object=company_ob, device__isnull=False,
                                                           device__device_type__isnull=False, deleted=False)
            if metering_points.count():
                for mp in metering_points:
                    metering_point_datas = MeteringPointData.objects.filter(metering_point=mp, report_type=report_type,
                                                                            device_type=mp.device.device_type, deleted=False)
                    if date_from:
                        metering_point_datas = metering_point_datas.filter(timestamp__gte=date_from)
                    if date_to:
                        metering_point_datas = metering_point_datas.filter(timestamp__lt=date_to)
                    metering_point_datas = metering_point_datas.order_by("timestamp")
                    rec['metering_points'].append({
                            'metering_point': mp,
                            'device': mp.device,
                            'headers': mp.device.device_type.parameters,
                            'values': metering_point_datas,
                    })

            result['company_objects'].append(rec)
        result["user"] = user if not report_user else report_user
        result["date"] = timezone.now()
        result["report_type"] = report_type
        result['date_from'] = date_from
        result['date_to'] = date_to
        return result

    metering_points = MeteringPoint.objects.filter(company_object__id__in = object_ids, device__isnull = False, device__device_type__isnull = False, deleted = False)
#     print("metering_points = ", metering_points)
    result = {
        "headers": [],
        "values": [],
    }
    device = None
    if metering_points.count():
        device = metering_points.first().device
        #FIXME придумать как быть, если у разных точек учета разный набор столбцов
        result["headers"] = metering_points.first().device.device_type.parameters
#         print(result["headers"])
    point_ids = metering_points.values_list("id", flat = True)
    metering_point_datas = MeteringPointData.objects.filter(metering_point__id__in = point_ids, report_type = report_type, deleted = False)
    if date_from:
        metering_point_datas = metering_point_datas.filter(timestamp__gte = date_from)            
    if date_to:
        metering_point_datas = metering_point_datas.filter(timestamp__lt = date_to)            
    metering_point_datas = metering_point_datas.order_by("timestamp")
    result["values"] = metering_point_datas
    result["headers_amount"] = len(result["headers"])
    result["user"] = user if not report_user else report_user
    result["company_objects"] = company_objects
    result["date"] = timezone.now()
    result["for_excel"] = for_excel
    result["device"] = device
    result["report_type"] = report_type
    return result


def generate_separate_objects_report(request, report_user = None, report_task = None):
    report_type = None
    date_from = None
    date_to = None
    object_ids = None
    company_objects = None
    user = None
    if request:
        report_type = request.GET.get("report_type", "DAILY")
        objects = request.GET.get("objects", "")
        date_from = request.GET.get("date_from", None)
        date_to = request.GET.get("date_to", None)
        if date_from:
            date_from = dateparse.parse_datetime(date_from)
        if date_to:
            date_to = dateparse.parse_datetime(date_to)
        object_ids = None
        if objects:
            object_ids = json.loads("[%s]" % objects)
    #         print("object_ids = " , object_ids)
        user = request.user
        permissions = request.user.get_permissions()
        company_objects = CompanyObject.objects.filter(deleted = False)
        if not permissions.status_admin_system:
            company_objects = company_objects.filter(company = user.employer)
        if object_ids:
            company_objects = company_objects.filter(id__in = object_ids)
    #         print("company_objects = ", company_objects)
        object_ids = company_objects.values_list("id", flat = True)
    else:
        report_type = report_task.report_type
        date_from = report_task.date_from
        date_to = report_task.date_to
        object_ids = []
        for object in report_task.company_objects.all():
            object_ids.append(object.id)
        company_objects = report_task.company_objects.all()
    result = []
    device_type = None
    #тут есть проблема, что для вывода отчета надо знать в какой шаблон его выводить. Шаблон задается в device type. но у нас отчет идет по списку объектов, в которых могут быть совершенно разнородные точки учета и приборы. Изначально надо было отчет делать по точкам учета. Поэтому тут костыль - берется и запоминается первый же найденный тип прибора, и в дальнейшем он используется для получения шаблона
    for company_object in company_objects:
        record = {
            "user": user if not report_user else report_user,
            "company_object": company_object,
            "date": timezone.now(),
            "report_type": report_type,
        }
        metering_points = MeteringPoint.objects.filter(company_object = company_object, device__isnull = False, device__device_type__isnull = False, deleted = False)
        if not device_type and metering_points.count():
            device_type = metering_points[0].device.device_type
        point_ids = metering_points.values_list("id", flat = True)
        metering_point_datas = MeteringPointData.objects.filter(metering_point__id__in = point_ids, report_type = report_type, deleted = False)
        if date_from:
            metering_point_datas = metering_point_datas.filter(timestamp__gte = date_from)
            record["date_from"] = date_from
        if date_to:
            metering_point_datas = metering_point_datas.filter(timestamp__lt = date_to)            
            record["date_to"] = date_to
        metering_point_datas = metering_point_datas.order_by("timestamp")
        record["rows"] = list(metering_point_datas)
        result.append(record)
    return result, device_type


def create_pdf_report(request, data, template_file_name, headers = None):
    avail_columns = "t1,t2,dt,P1,P2,dp,V1,V2,M1,M2,dM,tx,Px,Qо,BНP,BOC,НС".split(",")
    #имеется 18 колонок, надо их включить или выключить в зависимости от того, что пришло из фронта
    #дата всегда есть
    columns = [True]
    enabled_columns = 1
    fullwidth = ""
    if headers:
        header_values = headers.split(",")
        values = []
        for header_value in header_values:
            values.append(header_value.lower())
        for avail_column in avail_columns:
            ac = avail_column.lower()
            if ac in values:
                columns.append(True)
                enabled_columns += 1
            else:
                columns.append(False)
    else:
        for i in range(1, 18):
            columns.append(True)
            enabled_columns += 1
    print("enabled columns = ", enabled_columns)
    if enabled_columns > len(avail_columns) / 2:
        fullwidth = "fullwidth"
    context = {
        "columns": columns,
        "data": data,
        "fullwidth": fullwidth,
    }
    return PDFTemplateResponse(
        request, 
        template_file_name, 
        context, 
        filename = "report.pdf",
        cmd_options = {
            'encoding': 'utf8',
            'quiet': True,
            'orientation': 'landscape',
        },
    )


def create_docx_report(data, template_file_name, output_file_name, headers = None):

    def hide_column(table, column_index):
        table.columns[column_index].width = 0
        for cell in table.columns[column_index].cells:
            cell.text = ""
    
    DATA_TABLE_INDEX = 2
    header_values = None
    if headers:
        header_values = headers.split(",")
        values = []
        for header_value in header_values:
            values.append(header_value.lower())
        header_values = values        
    django_engine = engines['django']
    document = Document(template_file_name)
    #запомним, сколько изначально таблиц в документе
    tables_amount = len(document.tables)
    for i in range(1, len(data)):
        #добавляем страницы
        document.add_page_break()
        for table_index in range(0, tables_amount):
            template = document.tables[table_index]
            tbl = template._tbl
            # Here we do the copy of the table
            new_tbl = deepcopy(tbl)    
            paragraph = document.add_paragraph()
            # After that, we add the previously copied table
            paragraph._p.addnext(new_tbl)
    tables = document.tables
    table_index = 0
    for context in data:
        context_rows = context["rows"]
        context["rows"] = None
        for i in range(0, tables_amount):
            table = tables[table_index + i]
            columns = len(table.columns)
            if i == DATA_TABLE_INDEX:
                #это таблица с данными. там необходимо сделать столько строк, сколько строк с данными и в цикле считать туда все данные
                if header_values:
                    #уберем колонки, отсутствующие в данных
                    for column in range(1, columns):
                        value = table.cell(0, column).text
                        value = value.split(",")[0].split("\n")[0].strip().lower()
                        print("header = %s" % value)
                        if not value in header_values:
                            print("hidding %s" % value)
                            hide_column(table, column)
                if len(context_rows) > 1:
                    template_row = table.rows[1]
                    for row_index in range(1, len(context_rows)):
                        row = table.add_row()
                        for col_index in range(0, len(template_row.cells)):
                            template_cell = template_row.cells[col_index]
                            template_paragraph = template_cell.paragraphs[0]
                            cell = row.cells[col_index]
                            cell.text = template_paragraph.text
                            paragraph = cell.paragraphs[0]
                            paragraph.style = template_paragraph.style
                for row_index in range(0, len(context_rows)):
                    context_row = context_rows[row_index]
                    context["row"] = context_row
                    for column in range(0, columns):
                        cell = table.cell(1 + row_index, column)
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            template = django_engine.from_string("{% load call_method %}" + paragraph.text)
                            rendered_text = template.render(context)
                            if not paragraph.text == rendered_text:
                                paragraph.text = rendered_text
            #заполняем остальные поля
            rows = len(table.rows)
            for row in range(0, rows):
                for column in range(0, columns):
                    cell = table.cell(row, column)
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        template = django_engine.from_string("{% load call_method %}" + paragraph.text)
                        rendered_text = template.render(context)
                        if not paragraph.text == rendered_text:
                            paragraph.text = rendered_text
        table_index += tables_amount
    document.save(output_file_name)


def set_user_role(user):
    if user.department and user.employer:
        department = Department.objects.filter(name = user.department, company = user.employer, deleted = False)[:1]
#         department = Department.objects.filter(name = user.department, deleted = False)[:1]
#         print(department)
        if department.count():
            department = department[0]
            if department.department_type:
                department_type = DepartmentType.objects.filter(name = department.department_type, deleted = False)
                if not department_type.count():
                    try:
                        department_type = DepartmentType.objects.filter(id = int(department.department_type), deleted = False)
                    except Exception as ex:
                        traceback.print_exc()
#                 print(department_type)
                if department_type.count():
                    department_type = department_type[0]
                    if department_type.role:
                        user.role = department_type.role
#                         print(user.role)
                        user.save()
    

def check_mock_reports(objects, file_type, emails, report_type):
    return False   
#     print("check_mock_reports")
#     print(objects)
#     desired_report_type = 'HOURLY'
# #     desired_object_ids = [1]
#     desired_object_ids = [2, 9]
#     desired_file_type = ["DOCX", "PDF"]
#     if not report_type == desired_report_type or not file_type in desired_file_type:
#         print("report type or file type another")
#         return False
#     object_ids = json.loads("[%s]" % objects)    
#     object_ids.sort()
#     if not object_ids == desired_object_ids:
#         print("object ids different")
#         return False
#     for email in emails.split(","):
#         print("sending to ", email)
#         filename = None
#         content_type = None
#         if file_type == "DOCX":
#             filename = "%s/reports/email/report.docx" % settings.MEDIA_ROOT
#             content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         elif file_type == "PDF":
#             filename = "%s/reports/email/report.pdf" % settings.MEDIA_ROOT
#             content_type = "application/pdf"
#         msg = EmailMessage(
#             subject = "report",
#             body = '',
#             from_email = settings.DEFAULT_FROM_EMAIL,
#             to = [email],
#         )            
#         msg.attach_file(filename, content_type)
#         msg.send()
#     return True
    